from __future__ import annotations

import os
from pathlib import Path

from docx import Document


TEMPLATE_NAMES = ("内部结构计算书.docx", "核岛厂房计算书.docx")
CELL_INDEXES = (1, 2, 4, 5, 7, 9)


def patch_template(path: Path) -> None:
    document = Document(path)
    if not document.tables or len(document.tables[0].rows) <= 28:
        raise RuntimeError(f"template cover table layout is invalid: {path}")
    row = document.tables[0].rows[28]
    for placeholder_index, cell_index in enumerate(CELL_INDEXES, start=1):
        cell = row.cells[cell_index]
        expected = f"{{{{ cover_code_{placeholder_index} }}}}"
        current = cell.text.strip()
        if current not in {"", expected}:
            raise RuntimeError(
                f"cover code cell is not blank: {path.name} row=29 cell={cell_index}"
            )
        paragraph = cell.paragraphs[0]
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = expected
        else:
            paragraph.add_run(expected)
    temporary = path.with_name(f".{path.stem}.cover-code.tmp.docx")
    document.save(temporary)
    os.replace(temporary, path)


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    template_root = repo_root / "documents_bin" / "calculation_book"
    for name in TEMPLATE_NAMES:
        patch_template(template_root / name)


if __name__ == "__main__":
    main()

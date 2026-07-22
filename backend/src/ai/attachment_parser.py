from __future__ import annotations

import json
import struct
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass(frozen=True)
class AttachmentParseResult:
    kind: str
    media_type: str
    extracted_text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: tuple[str, ...] = ()


class AttachmentParseError(ValueError):
    def __init__(self, message: str, *, code: str = "attachment_parse_failed") -> None:
        super().__init__(message)
        self.code = code


def parse_attachment(
    path: Path | str,
    *,
    original_name: str,
    declared_media_type: str,
    max_chars: int,
    cad_workspace: Path | str | None = None,
    oda_converter: Any | None = None,
) -> AttachmentParseResult:
    source = Path(path)
    suffix = Path(original_name).suffix.lower()
    if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
        result = _parse_image(source, suffix=suffix)
    elif suffix in {".txt", ".md"}:
        result = _parse_text(source, suffix=suffix)
    elif suffix == ".pdf":
        result = _parse_pdf(source)
    elif suffix == ".docx":
        result = _parse_docx(source)
    elif suffix == ".xlsx":
        result = _parse_xlsx(source)
    elif suffix in {".dxf", ".dwg"}:
        result = _parse_drawing(
            source,
            suffix=suffix,
            cad_workspace=Path(cad_workspace) if cad_workspace is not None else None,
            oda_converter=oda_converter,
        )
    else:
        raise AttachmentParseError(
            f"unsupported attachment extension: {suffix or '<none>'}",
            code="attachment_type_not_allowed",
        )

    declared = declared_media_type.strip().lower()
    if declared.startswith("image/") and result.media_type != declared:
        aliases = {"image/jpg": "image/jpeg"}
        if aliases.get(declared, declared) != result.media_type:
            raise AttachmentParseError(
                "image MIME type does not match file signature",
                code="attachment_signature_mismatch",
            )
    return _truncate(result, max_chars=max_chars)


def _parse_text(path: Path, *, suffix: str) -> AttachmentParseResult:
    content = path.read_bytes()
    decoded: str | None = None
    selected_encoding = ""
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            decoded = content.decode(encoding)
        except UnicodeDecodeError:
            continue
        selected_encoding = encoding
        break
    if decoded is None:
        raise AttachmentParseError(
            "text attachment is not valid UTF-8 or GB18030",
            code="attachment_text_encoding_unsupported",
        )
    return AttachmentParseResult(
        kind="document",
        media_type="text/markdown" if suffix == ".md" else "text/plain",
        extracted_text=decoded,
        metadata={"encoding": selected_encoding},
    )


def _parse_pdf(path: Path) -> AttachmentParseResult:
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
        sections = [
            f"[Page {index}]\n{page.extract_text() or ''}".rstrip()
            for index, page in enumerate(reader.pages, start=1)
        ]
    except Exception as exc:
        raise AttachmentParseError(
            f"PDF extraction failed: {exc}",
            code="attachment_pdf_parse_failed",
        ) from exc
    return AttachmentParseResult(
        kind="document",
        media_type="application/pdf",
        extracted_text="\n\n".join(sections),
        metadata={"page_count": len(reader.pages)},
    )


def _parse_docx(path: Path) -> AttachmentParseResult:
    from docx import Document

    try:
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        sections: list[str] = []
        if paragraphs:
            sections.append("[Paragraphs]\n" + "\n".join(paragraphs))
        for index, table in enumerate(document.tables, start=1):
            rows = [
                " | ".join(cell.text for cell in row.cells)
                for row in table.rows
            ]
            sections.append(f"[Table {index}]\n" + "\n".join(rows))
    except Exception as exc:
        raise AttachmentParseError(
            f"DOCX extraction failed: {exc}",
            code="attachment_docx_parse_failed",
        ) from exc
    return AttachmentParseResult(
        kind="document",
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        extracted_text="\n\n".join(sections),
        metadata={
            "paragraph_count": len(paragraphs),
            "table_count": len(document.tables),
        },
    )


def _parse_xlsx(path: Path) -> AttachmentParseResult:
    from openpyxl import load_workbook

    workbook = None
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
        sections: list[str] = []
        nonempty_cells = 0
        for sheet in workbook.worksheets:
            lines: list[str] = []
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    nonempty_cells += 1
                    lines.append(f"{cell.coordinate}={_cell_text(cell.value)}")
            sections.append(f"[Sheet: {sheet.title}]\n" + "\n".join(lines))
        sheet_count = len(workbook.sheetnames)
    except Exception as exc:
        raise AttachmentParseError(
            f"XLSX extraction failed: {exc}",
            code="attachment_xlsx_parse_failed",
        ) from exc
    finally:
        if workbook is not None:
            workbook.close()
    return AttachmentParseResult(
        kind="document",
        media_type=(
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ),
        extracted_text="\n\n".join(sections),
        metadata={"sheet_count": sheet_count, "nonempty_cell_count": nonempty_cells},
    )


def _parse_image(path: Path, *, suffix: str) -> AttachmentParseResult:
    content = path.read_bytes()
    metadata: dict[str, Any] = {}
    if suffix == ".png":
        if len(content) < 24 or content[:8] != b"\x89PNG\r\n\x1a\n":
            raise AttachmentParseError(
                "PNG signature validation failed",
                code="attachment_signature_mismatch",
            )
        metadata["width"], metadata["height"] = struct.unpack(">II", content[16:24])
        media_type = "image/png"
    elif suffix in {".jpg", ".jpeg"}:
        if len(content) < 4 or not content.startswith(b"\xff\xd8\xff"):
            raise AttachmentParseError(
                "JPEG signature validation failed",
                code="attachment_signature_mismatch",
            )
        media_type = "image/jpeg"
    else:
        if len(content) < 12 or content[:4] != b"RIFF" or content[8:12] != b"WEBP":
            raise AttachmentParseError(
                "WebP signature validation failed",
                code="attachment_signature_mismatch",
            )
        media_type = "image/webp"
    return AttachmentParseResult(
        kind="image",
        media_type=media_type,
        extracted_text="",
        metadata=metadata,
    )


def _parse_drawing(
    path: Path,
    *,
    suffix: str,
    cad_workspace: Path | None,
    oda_converter: Any | None,
) -> AttachmentParseResult:
    from src.cad.ai.element_package_exporter import process_source

    workspace = cad_workspace or path.parent / "cad"
    drawing = process_source(
        source_path=path,
        dxf_dir=workspace / "dxf",
        oda=oda_converter,
        max_geometry_elements=500,
    )
    if drawing.get("status") != "ok":
        errors = drawing.get("errors") or []
        message = errors[0].get("message") if errors else "unknown CAD parse error"
        raise AttachmentParseError(
            f"CAD extraction failed: {message}",
            code="attachment_cad_parse_failed",
        )

    context = {
        "schema_version": "drawing-attachment-context@0.1",
        "source_name": drawing.get("source_name"),
        "source_format": suffix.removeprefix("."),
        "project_no_inferred": drawing.get("project_no_inferred"),
        "unit_no_inferred": drawing.get("unit_no_inferred"),
        "entity_type_counts": drawing.get("entity_type_counts", []),
        "layers": drawing.get("layers", []),
        "geometry_elements": drawing.get("geometry_elements", []),
        "geometry_elements_truncated": drawing.get(
            "geometry_elements_truncated",
            False,
        ),
        "text_elements": drawing.get("text_elements", []),
        "frames": drawing.get("frames", []),
        "semantic_summary": drawing.get("semantic_summary", {}),
    }
    entity_count = sum(
        int(item.get("count", 0))
        for item in drawing.get("entity_type_counts", [])
    )
    return AttachmentParseResult(
        kind="drawing",
        media_type="application/dxf" if suffix == ".dxf" else "application/acad",
        extracted_text=json.dumps(context, ensure_ascii=False, indent=2),
        metadata={
            "source_format": suffix.removeprefix("."),
            "conversion_status": drawing.get("conversion", {}).get("status"),
            "entity_count": entity_count,
            "frame_count": len(drawing.get("frames", [])),
            "text_element_count": len(drawing.get("text_elements", [])),
        },
    )


def _truncate(result: AttachmentParseResult, *, max_chars: int) -> AttachmentParseResult:
    limit = max(int(max_chars), 1)
    if len(result.extracted_text) <= limit:
        return result
    metadata = dict(result.metadata)
    metadata["truncated"] = True
    metadata["original_chars"] = len(result.extracted_text)
    return AttachmentParseResult(
        kind=result.kind,
        media_type=result.media_type,
        extracted_text=result.extracted_text[:limit],
        metadata=metadata,
        warnings=(*result.warnings, "content_truncated"),
    )


def _cell_text(value: Any) -> str:
    isoformat = getattr(value, "isoformat", None)
    if callable(isoformat):
        return str(isoformat())
    return str(value)

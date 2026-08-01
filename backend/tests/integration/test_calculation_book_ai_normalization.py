from __future__ import annotations

import json
import sys
import time
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from fastapi.testclient import TestClient as FastApiTestClient
from openpyxl import Workbook
from PIL import Image

from src.ai.chat_client import ChatCompletionResult
from src.ai.reinforcement_task_normalizer import (
    ReinforcementTaskNormalizer,
    ReinforcementTaskNormalizerLimits,
)
from src.calculation_book.executor import (
    CalculationBookJobExecutor,
    ReinforcementNormalizerMetadata,
)
from src.calculation_book.ocr import StressLegendReading
from src.calculation_book.processor import (
    CalculationBookAssets,
    CalculationBookProcessor,
)
from src.config import SpecLoader, reload_config


class TestClient(FastApiTestClient):
    def __enter__(self) -> TestClient:
        client = super().__enter__()
        response = client.post(
            "/api/auth/login",
            json={"account_id": "hbjjswd", "password": "password"},
        )
        assert response.status_code == 200, response.text
        client.headers["Authorization"] = f"Bearer {response.json()['token']}"
        return self


class FakeStructuredGateway:
    def __init__(self, payload: dict[str, object]) -> None:
        self.payload = payload
        self.calls: list[list[dict[str, Any]]] = []

    def complete(
        self,
        messages: list[dict[str, Any]],
        *,
        tools: list[dict[str, Any]] | None = None,
    ) -> ChatCompletionResult:
        assert tools is None
        self.calls.append(messages)
        return ChatCompletionResult(
            content=json.dumps(self.payload, ensure_ascii=False),
            raw_model="fake-structured-gateway",
        )


def _configure_api_env(monkeypatch: Any, tmp_path: Path) -> Path:
    repo_root = Path(__file__).resolve().parents[3]
    monkeypatch.setenv(
        "FANBAN_SPEC_PATH",
        str(repo_root / "documents" / "参数规范.yaml"),
    )
    monkeypatch.setenv(
        "FANBAN_RUNTIME_SPEC_PATH",
        str(repo_root / "documents" / "参数规范_运行期.yaml"),
    )
    monkeypatch.setenv("FANBAN_STORAGE_DIR", str(tmp_path / "storage"))
    monkeypatch.setenv("FANBAN_UPLOAD_LIMITS__MAX_FILES", "3")
    monkeypatch.setenv("FANBAN_UPLOAD_LIMITS__MAX_TOTAL_MB", "20")
    SpecLoader.clear_cache()
    reload_config()
    if str(repo_root) not in sys.path:
        sys.path.insert(0, str(repo_root))
    return repo_root


def _write_image(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    Image.new("RGB", (1200, 800), "white").save(path)


def _build_archive(
    tmp_path: Path,
    *,
    standard: bool,
    include_slab_figures: bool,
) -> Path:
    source = tmp_path / ("standard-source" if standard else "ai-source")
    for name in (
        "N5012-X.png",
        "N5012-Y.png",
        "N5012-Z.png",
        "01/layout.png",
        "02/model.png",
    ):
        _write_image(source / name)
    if include_slab_figures:
        for name in (
            "11.45-TOP-X.png",
            "11.45-BOTTOM-X.png",
            "11.45-TOP-Y.png",
            "11.45-BOTTOM-Y.png",
            "11.45-Z.png",
        ):
            _write_image(source / name)

    workbook = Workbook()
    sheet = workbook.active
    if standard:
        sheet.append(
            [
                "构件编号及位置",
                "单侧水平钢筋(对称配筋)",
                "单侧竖向钢筋(对称配筋)",
                "拉筋",
            ]
        )
        sheet.append(
            [
                "N5012 墙",
                "1D32间距200",
                "1D28间距200",
                "1C14间距400*400",
            ]
        )
    else:
        sheet.title = "墙体配筋"
        sheet.append(["墙号", "水平筋", "竖向筋", "拉筋"])
        sheet.append(
            [
                "N5012",
                "未知写法",
                "1 28@200",
                "1A14间距400*400",
            ]
        )
    workbook.save(source / "墙体配筋结果.xlsx")
    workbook.close()

    archive_path = tmp_path / ("standard.zip" if standard else "nonstandard.zip")
    with zipfile.ZipFile(archive_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in source.rglob("*"):
            if path.is_file():
                archive.write(path, path.relative_to(source).as_posix())
    return archive_path


def _params(*, include_slab_stress: bool) -> dict[str, object]:
    return {
        "template_type": "internal_structure",
        "project_no": "JQ",
        "project_name": "浙江金七门核电厂1、2号机组",
        "internal_code": "JQ00-NN-001",
        "version": "A",
        "subproject_code": "RX",
        "subproject_name": "内部结构",
        "design_phase": "施工图设计",
        "document_name": "0.000m~15.000m配筋计算书",
        "workshop_length": 72.5,
        "workshop_width": 48.0,
        "raft_slab_top_elevation": -8.5,
        "roof_top_elevation": 31.2,
        "factory_extreme_min_temperature": -18.0,
        "factory_extreme_max_temperature": 39.0,
        "site_soil_temperature": 15.0,
        "include_slab_stress": include_slab_stress,
    }


def _reading(direction: str) -> StressLegendReading:
    if direction == "Z":
        return StressLegendReading(
            smn=0,
            smx=0,
            legend_values=(),
            is_zero_result=True,
        )
    return StressLegendReading(
        smn=0,
        smx=800,
        legend_values=tuple(800 * index / 9 for index in range(10)),
    )


def _poll_job(
    client: TestClient,
    job_id: str,
    *,
    timeout_seconds: float = 20,
) -> dict[str, Any]:
    deadline = time.monotonic() + timeout_seconds
    while time.monotonic() < deadline:
        response = client.get(f"/api/jobs/{job_id}")
        assert response.status_code == 200, response.text
        payload = response.json()
        if payload["status"] in {"succeeded", "failed", "cancelled"}:
            return payload
        time.sleep(0.05)
    raise AssertionError(f"calculation job {job_id} did not finish")


def _document_text(path: Path) -> str:
    document = Document(path)
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return "\n".join(paragraph.text for paragraph in paragraphs)


def test_formal_task_flow_normalizes_once_leaves_partial_fields_blank_and_downloads(
    monkeypatch,
    tmp_path: Path,
) -> None:
    repo_root = _configure_api_env(monkeypatch, tmp_path)
    gateway = FakeStructuredGateway(
        {
            "schema_version": "1",
            "source_row_count": 1,
            "rows": [
                {
                    "kind": "wall",
                    "status": "needs_review",
                    "wall_id": "N5012",
                    "X": None,
                    "Y": "1D28间距200",
                    "Z": "1C14间距400*400",
                    "reason": "水平筋写法无法确定",
                    "blank_fields": ["X"],
                    "source_sheet": "墙体配筋",
                    "source_row": 2,
                    "source_cells": {
                        "wall": "A2",
                        "X": "B2",
                        "Y": "C2",
                        "Z": "D2",
                    },
                }
            ],
        }
    )
    normalizer = ReinforcementTaskNormalizer(
        client=gateway,
        skill_root=repo_root / "tools" / "ai" / "reinforcement-table-normalizer",
        limits=ReinforcementTaskNormalizerLimits(
            max_non_empty_cells=100,
            max_snapshot_chars=100_000,
            max_skill_chars=100_000,
        ),
    )
    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(
            template_root=repo_root / "documents_bin" / "calculation_book",
        ),
        ocr_recognizer=lambda _path, direction: _reading(direction),
    )
    executor = CalculationBookJobExecutor(
        processor=processor,
        normalizer=normalizer,
        normalizer_metadata=ReinforcementNormalizerMetadata(
            model="fake-structured",
            profile="integration",
        ),
    )

    import API.app.runtime as runtime_module
    from API.app.main import create_app
    from API.app.runtime import PipelineJobProcessor

    monkeypatch.setattr(
        runtime_module,
        "recognize_stress_legend",
        lambda _path, *, direction, **_kwargs: _reading(direction),
    )

    app = create_app(
        job_processor=PipelineJobProcessor(
            calculation_book_executor_factory=lambda: executor,
        )
    )
    nonstandard = _build_archive(
        tmp_path,
        standard=False,
        include_slab_figures=True,
    )
    standard = _build_archive(
        tmp_path,
        standard=True,
        include_slab_figures=False,
    )

    with TestClient(app) as client:
        preflight = client.post(
            "/api/jobs/calculation-books/preflight",
            data={"include_slab_stress": "true"},
            files={
                "archive": (
                    nonstandard.name,
                    nonstandard.read_bytes(),
                    "application/zip",
                )
            },
        )
        assert preflight.status_code == 200, preflight.text
        preflight_payload = preflight.json()
        assert preflight_payload["requires_ai_normalization"] is True
        assert preflight_payload["ai_confirmation_message"] == (
            "您上传的墙体配筋表非标准格式，程序将启动人工智能。"
        )
        assert preflight_payload["ai_reinforcement_expected_source_row_count"] == 1
        assert preflight_payload["slab_figure_count"] == 5
        assert gateway.calls == []

        unconfirmed = client.post(
            "/api/jobs/calculation-books",
            data={
                "params_json": json.dumps(
                    {
                        **_params(include_slab_stress=True),
                        "preflight_token": preflight_payload["preflight_token"],
                    },
                    ensure_ascii=False,
                )
            },
        )
        assert unconfirmed.status_code == 422
        assert gateway.calls == []

        created = client.post(
            "/api/jobs/calculation-books",
            data={
                "params_json": json.dumps(
                    {
                        **_params(include_slab_stress=True),
                        "preflight_token": preflight_payload["preflight_token"],
                        "confirm_ai_normalization": True,
                    },
                    ensure_ascii=False,
                )
            },
        )
        assert created.status_code == 201, created.text
        ai_job_id = created.json()["jobs"][0]["job_id"]
        ai_detail = _poll_job(client, ai_job_id)
        assert ai_detail["status"] == "succeeded", ai_detail
        assert len(gateway.calls) == 1
        prompt = "\n".join(
            str(message["content"])
            for message in gateway.calls[0]
        )
        assert "skill_id=reinforcement_table_normalizer" in prompt
        assert "source_row_count 与 rows 数量都必须等于 1" in prompt
        output = ai_detail["calculation_book_output"]
        assert output["ai_normalized"] is True
        assert output["ai_normalization"]["call_count"] == 1
        assert output["ai_normalization"]["source_row_count"] == 1
        assert output["warning_count"] == 2
        warnings = {item["code"]: item for item in output["warnings"]}
        assert warnings["needs_review"]["identity"] == "N5012"
        assert warnings["needs_review"]["blank_fields"] == ["X"]
        assert warnings["image_only_slab"]["identity"] == "11.45"
        assert warnings["image_only_slab"]["blank_fields"] == [
            "top_x",
            "bottom_x",
            "top_y",
            "bottom_y",
            "z",
        ]
        download = client.get(
            f"/api/jobs/{ai_job_id}/download/calculation-book"
        )
        assert download.status_code == 200
        assert download.content.startswith(b"PK")
        output_path = Path(
            client.app.state.runtime.job_manager.get_job(ai_job_id).artifacts.calculation_docx
        )
        text = _document_text(output_path)
        assert "N5012-水平向钢筋计算配筋面积" not in text
        assert "墙N5012-竖向钢筋计算配筋面积" in text
        assert "11.45m 楼板顶层水平配筋云图" in text
        assert "11.45m楼板顶层水平钢筋" not in text

        standard_preflight = client.post(
            "/api/jobs/calculation-books/preflight",
            data={"include_slab_stress": "false"},
            files={
                "archive": (
                    standard.name,
                    standard.read_bytes(),
                    "application/zip",
                )
            },
        )
        assert standard_preflight.status_code == 200, standard_preflight.text
        assert standard_preflight.json()["requires_ai_normalization"] is False
        standard_created = client.post(
            "/api/jobs/calculation-books",
            data={
                "params_json": json.dumps(
                    {
                        **_params(include_slab_stress=False),
                        "preflight_token": standard_preflight.json()[
                            "preflight_token"
                        ],
                    },
                    ensure_ascii=False,
                )
            },
        )
        assert standard_created.status_code == 201, standard_created.text
        standard_job_id = standard_created.json()["jobs"][0]["job_id"]
        standard_detail = _poll_job(client, standard_job_id)
        assert standard_detail["status"] == "succeeded", standard_detail
        assert standard_detail["calculation_book_output"]["ai_normalized"] is False
        assert len(gateway.calls) == 1

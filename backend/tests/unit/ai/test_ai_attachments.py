from __future__ import annotations

import hashlib
import sqlite3
import struct
from datetime import UTC, datetime, timedelta
from pathlib import Path

import pytest


def _stores(tmp_path: Path):
    from src.ai.attachment_store import AiAttachmentStore
    from src.ai.chat_store import AiChatStore

    chat_store = AiChatStore(tmp_path / "chat" / "fanban_ai_chat.sqlite3")
    chat_store.initialize()
    attachment_store = AiAttachmentStore(chat_store)
    return chat_store, attachment_store


def test_attachment_store_creates_lists_and_hashes_owner_isolated_file(
    tmp_path: Path,
) -> None:
    chat_store, attachment_store = _stores(tmp_path)
    owner_key = "ip:10.0.0.8"
    conversation = chat_store.create_conversation(owner_key=owner_key, title="附件")
    content = b"attachment-marker"

    created = attachment_store.create_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        original_name="sensitive-plan.txt",
        media_type="text/plain",
        content=content,
    )

    stored_path = attachment_store.resolve_path(created)
    assert created.sha256 == hashlib.sha256(content).hexdigest()
    assert created.size_bytes == len(content)
    assert created.status == "uploaded"
    assert stored_path.read_bytes() == content
    assert "10.0.0.8" not in str(stored_path)
    assert "sensitive-plan" not in str(stored_path)
    assert attachment_store.get_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        attachment_id=created.attachment_id,
    ) == created
    assert attachment_store.list_attachments(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
    ) == [created]


def test_attachment_store_denies_cross_owner_access_and_binding(tmp_path: Path) -> None:
    chat_store, attachment_store = _stores(tmp_path)
    owner_key = "ip:10.0.0.8"
    other_owner = "ip:10.0.0.9"
    conversation = chat_store.create_conversation(owner_key=owner_key, title="附件")
    message = chat_store.add_message(conversation.conversation_id, "user", "查看附件")
    created = attachment_store.create_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        original_name="plan.pdf",
        media_type="application/pdf",
        content=b"%PDF-test",
    )

    assert attachment_store.list_attachments(
        owner_key=other_owner,
        conversation_id=conversation.conversation_id,
    ) == []
    assert attachment_store.get_attachment(
        owner_key=other_owner,
        conversation_id=conversation.conversation_id,
        attachment_id=created.attachment_id,
    ) is None
    assert attachment_store.bind_to_message(
        owner_key=other_owner,
        conversation_id=conversation.conversation_id,
        attachment_id=created.attachment_id,
        message_id=message.message_id,
    ) is None
    assert attachment_store.delete_attachment(
        owner_key=other_owner,
        conversation_id=conversation.conversation_id,
        attachment_id=created.attachment_id,
    ) is False

    bound = attachment_store.bind_to_message(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        attachment_id=created.attachment_id,
        message_id=message.message_id,
    )
    assert bound is not None
    assert bound.message_id == message.message_id


def test_attachment_store_delete_removes_database_row_and_file(tmp_path: Path) -> None:
    chat_store, attachment_store = _stores(tmp_path)
    owner_key = "ip:10.0.0.8"
    conversation = chat_store.create_conversation(owner_key=owner_key, title="附件")
    created = attachment_store.create_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        original_name="plan.txt",
        media_type="text/plain",
        content=b"delete-me",
    )
    stored_path = attachment_store.resolve_path(created)

    assert attachment_store.delete_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        attachment_id=created.attachment_id,
    ) is True
    assert stored_path.exists() is False
    assert attachment_store.get_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        attachment_id=created.attachment_id,
    ) is None


def test_chat_store_clear_and_delete_remove_attachment_files(tmp_path: Path) -> None:
    chat_store, attachment_store = _stores(tmp_path)
    owner_key = "ip:10.0.0.8"
    cleared = chat_store.create_conversation(owner_key=owner_key, title="清空")
    cleared_attachment = attachment_store.create_attachment(
        owner_key=owner_key,
        conversation_id=cleared.conversation_id,
        original_name="clear.txt",
        media_type="text/plain",
        content=b"clear",
    )
    cleared_path = attachment_store.resolve_path(cleared_attachment)

    assert chat_store.clear_conversation(cleared.conversation_id, owner_key) is True
    assert cleared_path.exists() is False
    assert attachment_store.list_attachments(
        owner_key=owner_key,
        conversation_id=cleared.conversation_id,
    ) == []

    deleted = chat_store.create_conversation(owner_key=owner_key, title="删除")
    deleted_attachment = attachment_store.create_attachment(
        owner_key=owner_key,
        conversation_id=deleted.conversation_id,
        original_name="delete.txt",
        media_type="text/plain",
        content=b"delete",
    )
    deleted_path = attachment_store.resolve_path(deleted_attachment)

    assert chat_store.delete_conversation(deleted.conversation_id, owner_key) is True
    assert deleted_path.exists() is False


def test_attachment_retention_cleans_expired_files(tmp_path: Path) -> None:
    chat_store, attachment_store = _stores(tmp_path)
    owner_key = "ip:10.0.0.8"
    conversation = chat_store.create_conversation(owner_key=owner_key, title="过期")
    created = attachment_store.create_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        original_name="expired.txt",
        media_type="text/plain",
        content=b"expired",
    )
    stored_path = attachment_store.resolve_path(created)
    now = datetime(2026, 7, 22, tzinfo=UTC)
    expired_at = (now - timedelta(days=31)).isoformat()
    with sqlite3.connect(chat_store.db_path) as connection:
        connection.execute(
            "UPDATE ai_attachments SET created_at = ? WHERE attachment_id = ?",
            (expired_at, created.attachment_id),
        )

    assert attachment_store.purge_expired_unbound(retention_days=30, now=now) == 1
    assert stored_path.exists() is False

    conversation_attachment = attachment_store.create_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        original_name="conversation.txt",
        media_type="text/plain",
        content=b"conversation",
    )
    conversation_path = attachment_store.resolve_path(conversation_attachment)
    with sqlite3.connect(chat_store.db_path) as connection:
        connection.execute(
            "UPDATE ai_conversations SET updated_at = ? WHERE conversation_id = ?",
            (expired_at, conversation.conversation_id),
        )

    assert chat_store.purge_expired(retention_days=30, now=now) == 1
    assert conversation_path.exists() is False


def test_attachment_parser_reads_gb18030_text_and_truncates(tmp_path: Path) -> None:
    from src.ai.attachment_parser import parse_attachment

    source = tmp_path / "说明.txt"
    source.write_bytes("中文标记 AI-TXT-0711 后续内容".encode("gb18030"))

    result = parse_attachment(
        source,
        original_name=source.name,
        declared_media_type="text/plain",
        max_chars=14,
    )

    assert result.kind == "document"
    assert "AI-TXT" in result.extracted_text
    assert len(result.extracted_text) == 14
    assert result.metadata["encoding"] == "gb18030"
    assert result.metadata["truncated"] is True
    assert "content_truncated" in result.warnings


def test_attachment_parser_reads_pdf_docx_and_xlsx_boundaries(tmp_path: Path) -> None:
    import fitz
    from docx import Document
    from openpyxl import Workbook

    from src.ai.attachment_parser import parse_attachment

    pdf_path = tmp_path / "sample.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "AI-PDF-0711")
    pdf.save(pdf_path)
    pdf.close()

    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("AI-DOCX-0711")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "参数"
    table.cell(0, 1).text = "VALUE-0711"
    document.save(docx_path)

    xlsx_path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "输入参数"
    sheet["A1"] = "AI-XLSX-0711"
    sheet["B2"] = 42
    workbook.save(xlsx_path)
    workbook.close()

    pdf_result = parse_attachment(
        pdf_path,
        original_name=pdf_path.name,
        declared_media_type="application/pdf",
        max_chars=20_000,
    )
    docx_result = parse_attachment(
        docx_path,
        original_name=docx_path.name,
        declared_media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        max_chars=20_000,
    )
    xlsx_result = parse_attachment(
        xlsx_path,
        original_name=xlsx_path.name,
        declared_media_type=(
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ),
        max_chars=20_000,
    )

    assert "[Page 1]" in pdf_result.extracted_text
    assert "AI-PDF-0711" in pdf_result.extracted_text
    assert "[Paragraphs]" in docx_result.extracted_text
    assert "[Table 1]" in docx_result.extracted_text
    assert "VALUE-0711" in docx_result.extracted_text
    assert "[Sheet: 输入参数]" in xlsx_result.extracted_text
    assert "A1=AI-XLSX-0711" in xlsx_result.extracted_text
    assert "B2=42" in xlsx_result.extracted_text


def test_attachment_parser_validates_image_signature(tmp_path: Path) -> None:
    from src.ai.attachment_parser import AttachmentParseError, parse_attachment

    png_path = tmp_path / "pixel.png"
    png_path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + b"\x00\x00\x00\x0dIHDR"
        + struct.pack(">II", 1, 1)
        + b"\x08\x06\x00\x00\x00"
    )

    result = parse_attachment(
        png_path,
        original_name=png_path.name,
        declared_media_type="image/png",
        max_chars=20_000,
    )

    assert result.kind == "image"
    assert result.media_type == "image/png"
    assert result.metadata["width"] == 1
    assert result.metadata["height"] == 1
    assert result.extracted_text == ""

    invalid = tmp_path / "fake.png"
    invalid.write_bytes(b"not-a-png")
    with pytest.raises(AttachmentParseError, match="signature"):
        parse_attachment(
            invalid,
            original_name=invalid.name,
            declared_media_type="image/png",
            max_chars=20_000,
        )


def test_attachment_parser_rejects_unsupported_extension(tmp_path: Path) -> None:
    from src.ai.attachment_parser import AttachmentParseError, parse_attachment

    source = tmp_path / "archive.zip"
    source.write_bytes(b"PK")

    with pytest.raises(AttachmentParseError, match="unsupported"):
        parse_attachment(
            source,
            original_name=source.name,
            declared_media_type="application/zip",
            max_chars=20_000,
        )


def test_attachment_parser_builds_structured_dxf_context(tmp_path: Path) -> None:
    import ezdxf

    from src.ai.attachment_parser import parse_attachment

    dxf_path = tmp_path / "drawing.dxf"
    document = ezdxf.new("R2018")
    document.layers.add("AI_LAYER")
    modelspace = document.modelspace()
    modelspace.add_line((0, 0), (10, 5), dxfattribs={"layer": "AI_LAYER"})
    modelspace.add_text(
        "AI-DXF-0711",
        dxfattribs={"layer": "AI_LAYER", "insert": (1, 2)},
    )
    document.saveas(dxf_path)

    result = parse_attachment(
        dxf_path,
        original_name=dxf_path.name,
        declared_media_type="application/dxf",
        max_chars=40_000,
        cad_workspace=tmp_path / "cad-work",
    )

    assert result.kind == "drawing"
    assert result.metadata["source_format"] == "dxf"
    assert result.metadata["entity_count"] == 2
    assert '"name": "AI_LAYER"' in result.extracted_text
    assert "AI-DXF-0711" in result.extracted_text
    assert '"type": "LINE"' in result.extracted_text


def test_attachment_parser_converts_dwg_then_uses_same_dxf_pipeline(
    tmp_path: Path,
) -> None:
    import ezdxf

    from src.ai.attachment_parser import parse_attachment

    converted_dxf = tmp_path / "converted.dxf"
    document = ezdxf.new("R2018")
    document.modelspace().add_text(
        "AI-DWG-0711",
        dxfattribs={"insert": (1, 2)},
    )
    document.saveas(converted_dxf)
    dwg_path = tmp_path / "drawing.dwg"
    dwg_path.write_bytes(b"mock-dwg")

    class FakeOdaConverter:
        def __init__(self) -> None:
            self.calls: list[tuple[Path, Path]] = []

        def dwg_to_dxf(self, source: Path, output_dir: Path) -> Path:
            self.calls.append((source, output_dir))
            return converted_dxf

    converter = FakeOdaConverter()
    result = parse_attachment(
        dwg_path,
        original_name=dwg_path.name,
        declared_media_type="application/acad",
        max_chars=40_000,
        cad_workspace=tmp_path / "cad-work",
        oda_converter=converter,
    )

    assert converter.calls == [(dwg_path, tmp_path / "cad-work" / "dxf")]
    assert result.kind == "drawing"
    assert result.metadata["source_format"] == "dwg"
    assert result.metadata["conversion_status"] == "converted"
    assert "AI-DWG-0711" in result.extracted_text


def test_chat_service_sends_image_url_and_keeps_user_content_plain(tmp_path: Path) -> None:
    from src.ai.attachment_store import AiAttachmentStore
    from src.ai.chat_service import AiChatRuntimeConfig, AiChatService
    from src.ai.chat_store import AiChatStore

    class FakeClient:
        def __init__(self) -> None:
            self.messages = []

        def complete(self, messages, *, tools=None):
            self.messages = messages
            return type(
                "ChatResult",
                (),
                {"content": "看到了图片", "usage": {}, "tool_calls": ()},
            )()

    chat_store = AiChatStore(tmp_path / "chat.sqlite3")
    chat_store.initialize()
    attachment_store = AiAttachmentStore(chat_store)
    client = FakeClient()
    service = AiChatService(
        store=chat_store,
        attachment_store=attachment_store,
        client=client,
        runtime=AiChatRuntimeConfig(
            attachments={"allowed_extensions": [".png"]},
        ),
    )
    owner_key = "ip:10.0.0.8"
    conversation = service.create_conversation(owner_key, title="图片")
    png = (
        b"\x89PNG\r\n\x1a\n"
        + b"\x00\x00\x00\x0dIHDR"
        + struct.pack(">II", 1, 1)
        + b"\x08\x06\x00\x00\x00"
    )
    attachment = service.upload_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        original_name="现场.png",
        media_type="image/png",
        content=png,
    )

    exchange = service.send_message(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        content="请识别图片",
        attachment_ids=[attachment.attachment_id],
        agent_id=None,
        skill_ids=[],
        mcp_server_ids=[],
        account_id=None,
    )

    current_user_content = client.messages[-1]["content"]
    assert isinstance(current_user_content, list)
    assert current_user_content[0] == {"type": "text", "text": "请识别图片"}
    assert current_user_content[1]["type"] == "image_url"
    assert current_user_content[1]["image_url"]["url"].startswith(
        "data:image/png;base64,"
    )
    assert exchange.user_message.content == "请识别图片"
    assert exchange.user_message.metadata["attachments"][0]["original_name"] == "现场.png"
    bound = attachment_store.get_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        attachment_id=attachment.attachment_id,
    )
    assert bound is not None
    assert bound.message_id == exchange.user_message.message_id


def test_chat_service_sends_parsed_document_as_untrusted_evidence(tmp_path: Path) -> None:
    from src.ai.attachment_store import AiAttachmentStore
    from src.ai.chat_service import AiChatRuntimeConfig, AiChatService
    from src.ai.chat_store import AiChatStore

    class FakeClient:
        def __init__(self) -> None:
            self.messages = []

        def complete(self, messages, *, tools=None):
            self.messages = messages
            return type(
                "ChatResult",
                (),
                {"content": "已读取", "usage": {}, "tool_calls": ()},
            )()

    chat_store = AiChatStore(tmp_path / "chat.sqlite3")
    chat_store.initialize()
    attachment_store = AiAttachmentStore(chat_store)
    client = FakeClient()
    service = AiChatService(
        store=chat_store,
        attachment_store=attachment_store,
        client=client,
        runtime=AiChatRuntimeConfig(
            attachments={"allowed_extensions": [".txt"]},
        ),
    )
    owner_key = "ip:10.0.0.8"
    conversation = service.create_conversation(owner_key, title="文档")
    attachment = service.upload_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        original_name="说明.txt",
        media_type="text/plain",
        content="AI-DOC-EVIDENCE-0711".encode(),
    )

    exchange = service.send_message(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        content="",
        attachment_ids=[attachment.attachment_id],
        agent_id=None,
        skill_ids=[],
        mcp_server_ids=[],
        account_id=None,
    )

    current_content = client.messages[-1]["content"]
    assert isinstance(current_content, str)
    assert "<attachments_untrusted>" in current_content
    assert "AI-DOC-EVIDENCE-0711" in current_content
    assert exchange.user_message.content == ""

    follow_up = service.send_message(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        content="刚才的附件标记是什么？",
        attachment_ids=[],
        agent_id=None,
        skill_ids=[],
        mcp_server_ids=[],
        account_id=None,
    )
    historical_user_contents = [
        message["content"]
        for message in client.messages[:-1]
        if message["role"] == "user"
    ]
    assert any("AI-DOC-EVIDENCE-0711" in content for content in historical_user_contents)
    assert follow_up.memory["used_history_messages"] == 2


def test_chat_service_rejects_cross_owner_and_failed_attachments(tmp_path: Path) -> None:
    import pytest

    from src.ai.attachment_store import AiAttachmentStore
    from src.ai.chat_service import (
        AiChatRuntimeConfig,
        AiChatService,
        AiChatValidationError,
    )
    from src.ai.chat_store import AiChatStore

    chat_store = AiChatStore(tmp_path / "chat.sqlite3")
    chat_store.initialize()
    attachment_store = AiAttachmentStore(chat_store)
    service = AiChatService(
        store=chat_store,
        attachment_store=attachment_store,
        client=object(),
        runtime=AiChatRuntimeConfig(
            attachments={"allowed_extensions": [".txt"]},
        ),
    )
    owner_key = "ip:10.0.0.8"
    conversation = service.create_conversation(owner_key, title="隔离")
    attachment = service.upload_attachment(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        original_name="说明.txt",
        media_type="text/plain",
        content=b"private",
    )
    other_conversation = service.create_conversation(
        "ip:10.0.0.9",
        title="其他用户",
    )

    with pytest.raises(AiChatValidationError, match="attachment"):
        service.send_message(
            owner_key="ip:10.0.0.9",
            conversation_id=other_conversation.conversation_id,
            content="越权",
            attachment_ids=[attachment.attachment_id],
            agent_id=None,
            skill_ids=[],
            mcp_server_ids=[],
            account_id=None,
        )

    attachment_store.mark_failed(
        owner_key=owner_key,
        conversation_id=conversation.conversation_id,
        attachment_id=attachment.attachment_id,
        error_code="test_failure",
    )
    with pytest.raises(AiChatValidationError, match="not ready"):
        service.send_message(
            owner_key=owner_key,
            conversation_id=conversation.conversation_id,
            content="读取失败附件",
            attachment_ids=[attachment.attachment_id],
            agent_id=None,
            skill_ids=[],
            mcp_server_ids=[],
            account_id=None,
        )


def test_attachment_api_upload_list_and_owner_isolation(tmp_path: Path) -> None:
    from fastapi import FastAPI
    from fastapi.testclient import TestClient

    from API.app.routers.ai import router
    from src.ai.attachment_store import AiAttachmentStore
    from src.ai.chat_service import AiChatRuntimeConfig, AiChatService
    from src.ai.chat_store import AiChatStore

    chat_store = AiChatStore(tmp_path / "chat.sqlite3")
    chat_store.initialize()
    service = AiChatService(
        store=chat_store,
        attachment_store=AiAttachmentStore(chat_store),
        client=object(),
        runtime=AiChatRuntimeConfig(
            attachments={"allowed_extensions": [".txt"]},
        ),
    )
    owner_key = "ip:10.0.0.8"
    conversation = service.create_conversation(owner_key, title="API 附件")
    app = FastAPI()
    app.include_router(router)
    app.state.ai_chat_service = service

    with TestClient(app, client=("10.0.0.8", 50000)) as client:
        uploaded = client.post(
            f"/api/ai/conversations/{conversation.conversation_id}/attachments",
            files={"file": ("说明.txt", b"API-FILE-0711", "text/plain")},
        )
        assert uploaded.status_code == 201
        payload = uploaded.json()
        assert payload["status"] == "ready"
        assert payload["original_name"] == "说明.txt"
        attachment_id = payload["attachment_id"]

        listed = client.get(
            f"/api/ai/conversations/{conversation.conversation_id}/attachments"
        )
        assert listed.status_code == 200
        assert [item["attachment_id"] for item in listed.json()] == [attachment_id]

    with TestClient(app, client=("10.0.0.9", 50001)) as other_client:
        blocked_list = other_client.get(
            f"/api/ai/conversations/{conversation.conversation_id}/attachments"
        )
        assert blocked_list.status_code == 404
        blocked_delete = other_client.delete(
            f"/api/ai/conversations/{conversation.conversation_id}/attachments/{attachment_id}"
        )
        assert blocked_delete.status_code == 404

    with TestClient(app, client=("10.0.0.8", 50000)) as owner_client:
        deleted = owner_client.delete(
            f"/api/ai/conversations/{conversation.conversation_id}/attachments/{attachment_id}"
        )
        assert deleted.status_code == 200

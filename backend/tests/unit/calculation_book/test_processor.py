from __future__ import annotations

import math
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image

from src.calculation_book.models import CalculationBookParams
from src.calculation_book.ocr import OcrRecognitionError, StressLegendReading
from src.calculation_book.processor import (
    CalculationBookAssets,
    CalculationBookProcessor,
    CalculationBookStage,
)
from src.calculation_book.rebar_recommender import (
    RebarSuggestionInput,
    RebarSuggestionResult,
    SelectedRebarSuggestion,
)
from src.calculation_book.reinforcement_input import (
    NormalizedReinforcementRow,
    NormalizedSlabReinforcementRow,
    ReinforcementRowIssue,
    SlabReinforcementSchedule,
    build_reinforcement_schedule,
    parse_linear_rebar_cell,
    parse_rebar_cell,
)

ASSET_ROOT = Path(__file__).resolve().parents[4] / "documents_bin" / "calculation_book"


def _write_png(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    Image.new("RGB", (1200, 800), "white").save(path)


def _build_zip(
    tmp_path: Path,
    *,
    include_slab: bool = False,
    include_middle: bool = False,
    slab_elevations: tuple[str, ...] = ("11.45",),
    include_workbook: bool = True,
    wall_ids: tuple[str, ...] = ("N5012",),
    ignored_root_images: tuple[str, ...] = (),
) -> Path:
    source = tmp_path / "source"
    names = [
        *(f"{wall_id}-{direction}.png" for wall_id in wall_ids for direction in "XYZ"),
        *ignored_root_images,
        "01/layout.png",
        "02/model.png",
    ]
    if include_slab:
        for elevation in slab_elevations:
            names.extend(
                [
                    f"{elevation}-TOP-X.png",
                    f"{elevation}-BOTTOM-X.png",
                    f"{elevation}-TOP-Y.png",
                    f"{elevation}-BOTTOM-Y.png",
                    f"{elevation}-Z.png",
                ]
            )
    if include_middle:
        names.extend(
            f"{elevation}-MIDDLE-{direction}.png"
            for elevation in slab_elevations
            for direction in ("X", "Y")
        )
    for name in names:
        _write_png(source / name)
    if include_workbook:
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(
            [
                "构件编号及位置",
                "单侧水平钢筋(对称配筋)",
                "单侧竖向钢筋(对称配筋)",
                "拉筋",
            ]
        )
        sheet.append(
            ["N5012 墙", "1D32间距200", "1D28间距200", "1C14间距400*400"]
        )
        if include_slab:
            slab_sheet = workbook.create_sheet("楼板配筋")
            slab_sheet.append(
                [
                    "标高",
                    "顶层水平",
                    "顶层竖向",
                    "中层水平",
                    "中层竖向",
                    "底层水平",
                    "底层竖向",
                    "纵向拉筋",
                ]
            )
            slab_sheet.append(
                [
                    11.45,
                    "1D36@200",
                    "1D40@200",
                    "1D32@200" if include_middle else None,
                    "1D34@200" if include_middle else None,
                    "1D30@200",
                    "1D28@200",
                    "1D16@200",
                ]
            )
        workbook.save(source / "计算书模板文件.xlsx")
    archive_path = tmp_path / "input.zip"
    with zipfile.ZipFile(archive_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in source.rglob("*"):
            if not path.is_file():
                continue
            archive.write(path, path.relative_to(source).as_posix())
    return archive_path


def _params(
    *,
    include_slab_stress: bool = False,
    reinforcement_source: str = "provided",
    template_type: str = "internal_structure",
) -> CalculationBookParams:
    return CalculationBookParams.model_validate(
        {
            "template_type": template_type,
            "project_no": "JQ",
            "project_name": "浙江金七门核电厂1、2号机组",
            "internal_code": "JQ00-NN-001",
            "version": "A",
            "subproject_code": "RX",
            "subproject_name": "内部结构",
            "design_phase": "施工图设计",
            "document_name": "0.000m~15.000m配筋计算书",
            "workshop_length": 72.5,
            "workshop_width": 48.0,
            "raft_slab_top_elevation": -8.5,
            "roof_top_elevation": 31.2,
            "factory_extreme_min_temperature": -18.0,
            "factory_extreme_max_temperature": 39.0,
            "site_soil_temperature": 15.0,
            "include_slab_stress": include_slab_stress,
            "reinforcement_source": reinforcement_source,
        }
    )


def _all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _assert_mm2_uses_true_superscript(document: Document) -> None:
    found = 0
    for paragraph in _all_paragraphs(document):
        runs = paragraph.runs
        for index, run in enumerate(runs):
            for offset, character in enumerate(run.text):
                if character != "2":
                    continue
                prefix = "".join(item.text for item in runs[:index]) + run.text[:offset]
                if not prefix.endswith("mm"):
                    continue
                found += 1
                assert run.font.superscript is True
    assert found > 0


def test_processor_renders_a_real_docx_and_reports_all_stages(tmp_path: Path) -> None:
    stages: list[CalculationBookStage] = []

    def recognize(_path: Path, direction: str) -> StressLegendReading:
        if direction == "Z":
            return StressLegendReading(
                smn=0,
                smx=0,
                legend_values=(),
                is_zero_result=True,
            )
        return StressLegendReading(
            smn=0,
            smx=800,
            legend_values=tuple(800 * index / 9 for index in range(10)),
        )

    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(
            template_root=ASSET_ROOT,
        ),
        ocr_recognizer=recognize,
    )

    result = processor.process(
        archive_path=_build_zip(tmp_path),
        output_dir=tmp_path / "output",
        params=_params(),
        progress=lambda stage, _percent, _message, _details: stages.append(stage),
    )

    assert result.output_path.is_file()
    assert result.output_path.suffix == ".docx"
    assert result.figure_count == 3
    assert len(result.selections) == 3
    assert result.selections[2].actual_area == pytest.approx(
        math.pi * 7**2 * 2.5 * 2.5
    )
    assert stages == [
        CalculationBookStage.VALIDATE_ARCHIVE,
        CalculationBookStage.OCR_REINFORCEMENT,
        CalculationBookStage.SELECT_REBAR,
        CalculationBookStage.RENDER_CALCULATION_BOOK,
        CalculationBookStage.FINALIZE_ARTIFACT,
    ]

    document = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert (
        "墙N5012-拉筋钢筋计算配筋面积为0mm2/m。"
        "选用钢筋1排14@400x400（配筋面积为962.1 mm2/m）作为构造钢筋。 "
        "配筋结果包络计算结果。"
    ) in text
    _assert_mm2_uses_true_superscript(document)


def test_processor_inserts_five_slab_groups_before_wall_results(
    tmp_path: Path,
) -> None:
    def recognize(_path: Path, direction: str) -> StressLegendReading:
        if direction == "Z":
            return StressLegendReading(
                smn=0,
                smx=0,
                legend_values=(),
                is_zero_result=True,
            )
        return StressLegendReading(
            smn=0,
            smx=800,
            legend_values=tuple(800 * index / 9 for index in range(10)),
        )

    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(template_root=ASSET_ROOT),
        ocr_recognizer=recognize,
    )

    result = processor.process(
        archive_path=_build_zip(tmp_path, include_slab=True),
        output_dir=tmp_path / "output",
        params=_params(include_slab_stress=True),
    )

    assert result.figure_count == 8
    document = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in _all_paragraphs(document))
    transition = (
        "墙体的配筋计算结果如下。"
        "配筋结果为单侧配筋量、其单位为mm2/m。"
    )
    assert text.index("11.45m楼板顶层水平钢筋") < text.index(transition)
    assert text.index(transition) < text.index("墙N5012-水平向钢筋")
    assert "11.45m楼板纵向拉筋计算配筋面积为0mm2/m。" in text
    assert text.count("11.45m楼板") == 5
    _assert_mm2_uses_true_superscript(document)


def _validated_override(*, include_slab: bool):
    wall_row = NormalizedReinforcementRow(
        wall_id="N5012",
        x=parse_linear_rebar_cell("1D32@200"),
        y=parse_linear_rebar_cell("1D28@200"),
        z=parse_rebar_cell("1C14@400*400", direction="Z"),
        source_sheet="AI",
        source_row=2,
        source_cells={"wall": "A2", "X": "B2", "Y": "C2", "Z": "D2"},
    )
    slab_schedule = None
    if include_slab:
        slab_schedule = SlabReinforcementSchedule(
            rows=(
                NormalizedSlabReinforcementRow(
                    elevation="11.45",
                    top_x=parse_linear_rebar_cell("1D36@200"),
                    top_y=parse_linear_rebar_cell("1D40@200"),
                    middle_x=None,
                    middle_y=None,
                    bottom_x=parse_linear_rebar_cell("1D30@200"),
                    bottom_y=parse_linear_rebar_cell("1D28@200"),
                    z=parse_rebar_cell("1C16@200*200", direction="Z"),
                    source_sheet="AI-slab",
                    source_row=2,
                    source_cells={
                        "elevation": "A2",
                        "top_x": "B2",
                        "top_y": "C2",
                        "bottom_x": "F2",
                        "bottom_y": "G2",
                        "z": "H2",
                    },
                ),
            )
        )
    return SimpleNamespace(
        wall_schedule=build_reinforcement_schedule(
            rows=(wall_row,),
            source_row_count=1,
            normalization_triggered=True,
        ),
        slab_schedule=slab_schedule,
        warnings=(),
        source_row_count=1 + int(include_slab),
    )


def _slab_review_warning(*, elevation: str):
    return SimpleNamespace(
        code="needs_review",
        scope="slab",
        identity=elevation,
        direction=None,
        source_sheet="AI-slab",
        source_row=3,
        source_cells={"elevation": "A3", "top_x": "B3"},
        original_values={"elevation": elevation, "top_x": "ambiguous"},
        resolved_values={
            "elevation": elevation,
            "top_y": "1D40间距200",
            "bottom_x": "1D30间距200",
            "bottom_y": "1D28间距200",
            "z": "1D16间距200",
        },
        reason="slab reinforcement needs review",
        blank_fields=("top_x",),
    )


def test_processor_uses_ai_schedule_override_after_one_safe_extraction(
    monkeypatch,
    tmp_path: Path,
) -> None:
    import src.calculation_book.processor as processor_module

    archive_path = _build_zip(tmp_path, include_slab=True)
    extraction_calls = 0
    real_extract = processor_module.validate_and_extract_archive

    def count_extract(*args, **kwargs):
        nonlocal extraction_calls
        extraction_calls += 1
        return real_extract(*args, **kwargs)

    monkeypatch.setattr(
        processor_module,
        "validate_and_extract_archive",
        count_extract,
    )
    monkeypatch.setattr(
        processor_module,
        "load_reinforcement_schedule",
        lambda *_args, **_kwargs: pytest.fail(
            "AI override must bypass deterministic wall workbook loading"
        ),
    )
    monkeypatch.setattr(
        processor_module,
        "load_slab_reinforcement_schedule",
        lambda *_args, **_kwargs: pytest.fail(
            "AI override must bypass deterministic slab workbook loading"
        ),
    )
    callback_calls: list[tuple[Path, bool]] = []

    def normalize(workbook_path: Path, include_slab: bool):
        assert workbook_path.is_file()
        callback_calls.append((workbook_path, include_slab))
        return _validated_override(include_slab=True)

    def recognize(_path: Path, direction: str) -> StressLegendReading:
        return StressLegendReading(
            smn=0,
            smx=0 if direction == "Z" else 800,
            legend_values=() if direction == "Z" else tuple(range(10)),
            is_zero_result=direction == "Z",
        )

    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(template_root=ASSET_ROOT),
        ocr_recognizer=recognize,
    )

    result = processor.process(
        archive_path=archive_path,
        output_dir=tmp_path / "override-output",
        params=_params(include_slab_stress=True),
        reinforcement_normalizer=normalize,
    )

    assert extraction_calls == 1
    assert len(callback_calls) == 1
    assert callback_calls[0][1] is True
    assert result.figure_count == 8
    assert result.normalization_warnings == ()


def test_processor_does_not_require_ai_slab_schedule_when_slab_is_disabled(
    monkeypatch,
    tmp_path: Path,
) -> None:
    import src.calculation_book.processor as processor_module

    monkeypatch.setattr(
        processor_module,
        "load_reinforcement_schedule",
        lambda *_args, **_kwargs: pytest.fail("AI override must be used"),
    )
    monkeypatch.setattr(
        processor_module,
        "load_slab_reinforcement_schedule",
        lambda *_args, **_kwargs: pytest.fail("slab is disabled"),
    )
    seen: list[bool] = []

    def normalize(_workbook_path: Path, include_slab: bool):
        seen.append(include_slab)
        return _validated_override(include_slab=False)

    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(template_root=ASSET_ROOT),
        ocr_recognizer=lambda _path, direction: StressLegendReading(
            smn=0,
            smx=0 if direction == "Z" else 800,
            legend_values=() if direction == "Z" else tuple(range(10)),
            is_zero_result=direction == "Z",
        ),
    )

    result = processor.process(
        archive_path=_build_zip(tmp_path, include_slab=True),
        output_dir=tmp_path / "wall-only-output",
        params=_params(include_slab_stress=False),
        reinforcement_normalizer=normalize,
    )

    assert seen == [False]
    assert result.figure_count == 3


def test_ai_review_only_slab_rows_do_not_fail_calculation_generation(
    monkeypatch,
    tmp_path: Path,
) -> None:
    import src.calculation_book.processor as processor_module

    monkeypatch.setattr(
        processor_module,
        "load_reinforcement_schedule",
        lambda *_args, **_kwargs: pytest.fail("AI override must be used"),
    )
    warning = _slab_review_warning(elevation="11.45")
    validated = _validated_override(include_slab=False)
    validated.warnings = (warning,)

    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(template_root=ASSET_ROOT),
        ocr_recognizer=lambda _path, direction: StressLegendReading(
            smn=0,
            smx=0 if direction == "Z" else 800,
            legend_values=() if direction == "Z" else tuple(range(10)),
            is_zero_result=direction == "Z",
        ),
    )

    result = processor.process(
        archive_path=_build_zip(tmp_path, include_slab=True),
        output_dir=tmp_path / "review-only-slab-output",
        params=_params(include_slab_stress=True),
        reinforcement_normalizer=lambda _path, _include_slab: validated,
    )

    assert result.output_path.is_file()
    assert result.figure_count == 8
    assert result.normalization_warnings == (warning,)
    document = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in _all_paragraphs(document))
    assert "11.45m 楼板顶层水平配筋云图" in text
    assert "11.45m楼板顶层水平钢筋" not in text
    assert "11.45m楼板底层水平钢筋" in text


@pytest.mark.parametrize(
    ("include_middle", "expected_keys"),
    [
        (
            False,
            ("top_x", "bottom_x", "top_y", "bottom_y", "z"),
        ),
        (
            True,
            (
                "top_x",
                "middle_x",
                "bottom_x",
                "top_y",
                "middle_y",
                "bottom_y",
                "z",
            ),
        ),
    ],
)
def test_ai_missing_slab_rows_keeps_named_groups_blank_and_completes(
    monkeypatch,
    tmp_path: Path,
    include_middle: bool,
    expected_keys: tuple[str, ...],
) -> None:
    import src.calculation_book.processor as processor_module

    monkeypatch.setattr(
        processor_module,
        "load_reinforcement_schedule",
        lambda *_args, **_kwargs: pytest.fail("AI override must be used"),
    )
    validated = _validated_override(include_slab=False)
    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(template_root=ASSET_ROOT),
        ocr_recognizer=lambda _path, direction: StressLegendReading(
            smn=0,
            smx=0 if direction == "Z" else 800,
            legend_values=() if direction == "Z" else tuple(range(10)),
            is_zero_result=direction == "Z",
        ),
    )

    result = processor.process(
        archive_path=_build_zip(
            tmp_path,
            include_slab=True,
            include_middle=include_middle,
        ),
        output_dir=tmp_path / f"missing-slab-rows-{include_middle}",
        params=_params(include_slab_stress=True),
        reinforcement_normalizer=lambda _path, _include_slab: validated,
    )

    assert result.output_path.is_file()
    [warning] = [
        item for item in result.normalization_warnings
        if item.code == "image_only_slab"
    ]
    assert warning.identity == "11.45"
    assert warning.blank_fields == expected_keys
    document = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in _all_paragraphs(document))
    assert "11.45m 楼板顶层水平配筋云图" in text
    assert "11.45m楼板顶层水平钢筋" not in text
    if include_middle:
        assert "11.45m 楼板中层水平配筋云图" in text


def test_ai_partial_slab_schedule_matches_only_resolved_elevations(
    monkeypatch,
    tmp_path: Path,
) -> None:
    import src.calculation_book.processor as processor_module

    monkeypatch.setattr(
        processor_module,
        "load_reinforcement_schedule",
        lambda *_args, **_kwargs: pytest.fail("AI override must be used"),
    )
    warning = _slab_review_warning(elevation="15.95")
    validated = _validated_override(include_slab=True)
    validated.warnings = (warning,)
    validated.source_row_count = 3

    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(template_root=ASSET_ROOT),
        ocr_recognizer=lambda _path, direction: StressLegendReading(
            smn=0,
            smx=0 if direction == "Z" else 800,
            legend_values=() if direction == "Z" else tuple(range(10)),
            is_zero_result=direction == "Z",
        ),
    )

    result = processor.process(
        archive_path=_build_zip(
            tmp_path,
            include_slab=True,
            slab_elevations=("11.45", "15.95"),
        ),
        output_dir=tmp_path / "partial-slab-output",
        params=_params(include_slab_stress=True),
        reinforcement_normalizer=lambda _path, _include_slab: validated,
    )

    assert result.output_path.is_file()
    assert result.figure_count == 13
    assert result.normalization_warnings == (warning,)
    document = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "11.45m楼板顶层水平钢筋" in text


def test_ai_partial_wall_keeps_heading_and_image_but_blanks_only_x_values(
    monkeypatch,
    tmp_path: Path,
) -> None:
    import src.calculation_book.processor as processor_module

    monkeypatch.setattr(
        processor_module,
        "load_reinforcement_schedule",
        lambda *_args, **_kwargs: pytest.fail("AI override must be used"),
    )
    issue = ReinforcementRowIssue(
        source_sheet="AI",
        source_row=7,
        source_cells={"wall": "A7", "X": "B7", "Y": "C7", "Z": "D7"},
        original_values={"wall": "N5012", "X": "?"},
        original_wall_text="N5012",
        wall_id="N5012",
        error="X 向待确认",
    )
    warning = SimpleNamespace(
        code="needs_review",
        scope="wall",
        identity="N5012",
        direction="X",
        source_sheet="AI",
        source_row=7,
        source_cells=issue.source_cells,
        original_values=issue.original_values,
        resolved_values={
            "wall_id": "N5012",
            "Y": "1D28间距200",
            "Z": "1C14间距400*400",
        },
        reason="X 向待确认",
        blank_fields=("X",),
    )
    validated = SimpleNamespace(
        wall_schedule=build_reinforcement_schedule(
            rows=(),
            issues=(issue,),
            source_row_count=1,
            normalization_triggered=True,
        ),
        slab_schedule=None,
        warnings=(warning,),
        source_row_count=1,
    )
    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(template_root=ASSET_ROOT),
        ocr_recognizer=lambda _path, direction: StressLegendReading(
            smn=0,
            smx=0 if direction == "Z" else 800,
            legend_values=() if direction == "Z" else tuple(range(10)),
            is_zero_result=direction == "Z",
        ),
    )

    result = processor.process(
        archive_path=_build_zip(tmp_path),
        output_dir=tmp_path / "partial-wall-output",
        params=_params(),
        reinforcement_normalizer=lambda _path, _include_slab: validated,
    )

    assert [item.direction for item in result.selections] == ["Y", "Z"]
    assert result.warnings == (warning,)
    document = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "N5012-水平向" in text
    assert "墙N5012-水平向钢筋计算配筋面积" not in text
    assert "墙N5012-竖向钢筋计算配筋面积" in text
    assert len(document.inline_shapes) >= 3


def _select_first_suggestions(
    items: tuple[RebarSuggestionInput, ...],
) -> RebarSuggestionResult:
    selected: list[SelectedRebarSuggestion] = []
    for item in items:
        assert item.candidates, item.item_id
        candidate = item.candidates[0]
        selected.append(
            SelectedRebarSuggestion(
                item_id=item.item_id,
                member_kind=item.member_kind,
                member_id=item.member_id,
                direction=item.direction,
                smx=item.smx,
                target_area=item.target_area,
                candidate=candidate,
                configuration=parse_rebar_cell(
                    candidate.canonical_specification,
                    direction=item.direction,
                ).selected,
                reason="测试选择精确候选",
                source=(
                    "fixed_rule"
                    if item.direction == "Z" and item.smx == 0
                    else "ai"
                ),
            )
        )
    return RebarSuggestionResult(
        selected=tuple(selected),
        warnings=(),
        call_count=1,
        repair_round_count=0,
        skill_id="recommend-rebar-from-smx",
        skill_version="1.0.0",
        skill_sha256="sha-test",
        model="test-model",
    )


@pytest.mark.parametrize(
    "template_type",
    ["internal_structure", "nuclear_island_plant"],
)
def test_ai_suggested_processor_skips_excel_and_renders_five_or_seven_slab_items(
    tmp_path: Path,
    template_type: str,
) -> None:
    stages: list[CalculationBookStage] = []
    captured_items: list[RebarSuggestionInput] = []
    audit_events: list[tuple[str, dict[str, object]]] = []

    def recognize(_path: Path, direction: str) -> StressLegendReading:
        if direction == "Z":
            return StressLegendReading(0, 0, (), is_zero_result=True)
        return StressLegendReading(
            smn=0,
            smx=800,
            legend_values=tuple(800 * index / 9 for index in range(10)),
        )

    def suggest(items: tuple[RebarSuggestionInput, ...]) -> RebarSuggestionResult:
        captured_items.extend(items)
        return _select_first_suggestions(items)

    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(template_root=ASSET_ROOT),
        ocr_recognizer=recognize,
    )
    result = processor.process(
        archive_path=_build_zip(
            tmp_path,
            include_slab=True,
            include_middle=True,
            include_workbook=False,
        ),
        output_dir=tmp_path / "ai-output",
        params=_params(
            include_slab_stress=True,
            reinforcement_source="ai_suggested",
            template_type=template_type,
        ),
        reinforcement_normalizer=lambda *_args: pytest.fail(
            "无实配钢筋模式不得启动 Excel 规范化"
        ),
        rebar_suggester=suggest,
        progress=lambda stage, _percent, _message, _details: stages.append(stage),
        audit=lambda event, payload: audit_events.append((event, payload)),
    )

    assert stages == [
        CalculationBookStage.VALIDATE_ARCHIVE,
        CalculationBookStage.OCR_REINFORCEMENT,
        CalculationBookStage.AI_REBAR_SUGGESTION,
        CalculationBookStage.SELECT_REBAR,
        CalculationBookStage.RENDER_CALCULATION_BOOK,
        CalculationBookStage.FINALIZE_ARTIFACT,
    ]
    assert len(captured_items) == 10
    assert {item.item_id for item in captured_items if item.member_kind == "slab"} == {
        "slab:11.45:top_x",
        "slab:11.45:middle_x",
        "slab:11.45:bottom_x",
        "slab:11.45:top_y",
        "slab:11.45:middle_y",
        "slab:11.45:bottom_y",
        "slab:11.45:z",
    }
    assert result.ai_rebar_suggestion is not None
    assert result.ai_suggested_direction_count == 10
    assert result.ai_blank_direction_count == 0
    assert result.figure_count == 10
    assert [event for event, _payload in audit_events].count("archive_validated") == 1
    assert [event for event, _payload in audit_events].count("image_grouped") == 10
    assert [event for event, _payload in audit_events].count("ocr_completed") == 10
    assert [event for event, _payload in audit_events].count("word_entry_written") == 10
    document = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in _all_paragraphs(document))
    disclosure = (
        "以下配筋建议由人工智能根据结果云图 SMX 值并保留不低于 10% 的面积裕度生成，"
        "供设计人员复核。"
    )
    transition = (
        "墙体的配筋计算结果如下。"
        "配筋结果为单侧配筋量、其单位为mm2/m。"
    )
    assert disclosure in text
    assert "建议选用钢筋" in text
    assert text.index("11.45m楼板中层竖向钢筋") < text.index(transition)
    assert text.index(transition) < text.index("墙N5012-水平向钢筋")
    _assert_mm2_uses_true_superscript(document)


def test_ai_processor_isolates_ocr_failure_split_groups_and_unknown_images(
    tmp_path: Path,
) -> None:
    recognized_names: list[str] = []
    captured_ids: list[str] = []

    def recognize(path: Path, direction: str) -> StressLegendReading:
        recognized_names.append(path.name)
        if path.name == "S7157A-X.png":
            raise OcrRecognitionError("SMX 未识别")
        if direction == "Z":
            return StressLegendReading(0, 0, (), is_zero_result=True)
        return StressLegendReading(
            smn=0,
            smx=800,
            legend_values=tuple(800 * index / 9 for index in range(10)),
        )

    def suggest(items: tuple[RebarSuggestionInput, ...]) -> RebarSuggestionResult:
        captured_ids.extend(item.item_id for item in items)
        return _select_first_suggestions(items)

    processor = CalculationBookProcessor(
        assets=CalculationBookAssets(template_root=ASSET_ROOT),
        ocr_recognizer=recognize,
    )
    result = processor.process(
        archive_path=_build_zip(
            tmp_path,
            include_workbook=False,
            wall_ids=("S7157", "S7157A", "S7157-1", "S7157-2"),
            ignored_root_images=("无法识别名称.png",),
        ),
        output_dir=tmp_path / "ai-partial-output",
        params=_params(reinforcement_source="ai_suggested"),
        rebar_suggester=suggest,
    )

    assert all("S7157-1" not in item_id and "S7157-2" not in item_id for item_id in captured_ids)
    assert not any("S7157-1" in name or "S7157-2" in name for name in recognized_names)
    assert set(captured_ids) == {
        "wall:S7157:X",
        "wall:S7157:Y",
        "wall:S7157:Z",
        "wall:S7157A:Y",
        "wall:S7157A:Z",
    }
    assert result.figure_count == 13
    assert result.ai_suggested_direction_count == 5
    assert result.ai_blank_direction_count == 7
    warning_codes = {warning.code for warning in result.warnings}
    assert {
        "OCR_RECOGNITION_FAILED",
        "split_image_group",
        "UNKNOWN_IMAGE_NAME",
    } <= warning_codes
    document = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in _all_paragraphs(document))
    assert "S7157A-水平向" in text
    assert "墙S7157A-水平向钢筋计算配筋面积" not in text
    assert "S7157-1-水平向" in text
    assert "S7157-2-拉筋" in text
    assert "无法识别的应力图：无法识别名称" in text
